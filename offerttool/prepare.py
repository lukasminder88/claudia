"""Vorlage einmalig mit Ankern präparieren.

Die gelieferten Vorlagen (``Offerte_deCH_Miete.docx``, ``Offerte_deCH_Kauf.docx``)
sind ausgefüllte Beispieldokumente ohne Anker.  Dieser Schritt erzeugt daraus
**einmalig** die ankerbasierte Vorlage nach Abschnitt 3.

Die unvermeidbare Strukturerkennung passiert damit offline, überprüfbar und
genau einmal – nicht bei jeder Generierung.  Danach adressiert der Generator
ausschliesslich über ``w:tag``.  Das Ergebnis wird mit ``validate_template``
gegen den Ankerkatalog geprüft.

Der Aufbau folgt Abschnitt 9: die Kapitel 1.3 (Total) und die Musterzeilen der
Tabellen werden hier hergestellt, damit der Renderer nur noch klont.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.oxml.ns import qn

from .anchors import ALL_TAGS, BY_TAG, SWITCH_VARIANTS
from .docxutil.anchor_ops import tag_map
from .docxutil.tables import cells, rows, set_cell_paragraphs, set_tbl_look
from .docxutil.xmlutil import (
    W,
    clone,
    delete,
    make_paragraph,
    make_sdt,
    paragraph_text,
    sdt_content,
    wrap_in_sdt,
)
from .errors import OfferteError


class Locator:
    """Findet Blockelemente der Rohvorlage über ihre Struktur.

    Jede Suche ist eine Behauptung: trifft sie nicht genau einmal zu, bricht die
    Präparation ab.  So schlägt eine geänderte Vorlage hier auf – nicht später
    still im fertigen Dokument.
    """

    def __init__(self, doc) -> None:
        self.doc = doc
        self.body = doc.element.body

    def blocks(self) -> list:
        return [c for c in self.body.iterchildren() if c.tag in (W("w:p"), W("w:tbl"))]

    def paragraphs(self) -> list:
        return [c for c in self.body.iterchildren() if c.tag == W("w:p")]

    def tables(self) -> list:
        return self.body.findall(W("w:tbl"))

    def table_by_style(self, style_id: str, nth: int = 0):
        hits = [t for t in self.tables() if _tbl_style(t) == style_id]
        if len(hits) <= nth:
            raise OfferteError(
                "E101", f"Vorlage: keine {nth + 1}. Tabelle mit Style {style_id!r}"
            )
        return hits[nth]

    def para(self, style_id: str, startswith: str = "", nth: int = 0):
        hits = [
            p
            for p in self.paragraphs()
            if _p_style(p) == style_id and paragraph_text(p).strip().startswith(startswith)
        ]
        if len(hits) <= nth:
            raise OfferteError(
                "E101",
                f"Vorlage: kein {nth + 1}. Absatz [{style_id}] mit Anfang {startswith!r}",
            )
        return hits[nth]


def _tbl_style(tbl_el) -> str | None:
    pr = tbl_el.find(W("w:tblPr"))
    if pr is None:
        return None
    st = pr.find(W("w:tblStyle"))
    return st.get(qn("w:val")) if st is not None else None


def _p_style(p_el) -> str:
    ppr = p_el.find(W("w:pPr"))
    if ppr is None:
        return "Normal"
    st = ppr.find(W("w:pStyle"))
    return st.get(qn("w:val")) if st is not None else "Normal"


def _cell(tbl_el, row: int, col: int):
    tr = rows(tbl_el)[row]
    return cells(tr)[col]


def _cell_blocks(tc_el) -> list:
    return [c for c in tc_el.iterchildren() if c.tag in (W("w:p"), W("w:tbl"), W("w:sdt"))]


def _keep_rows(tbl_el, keep: list[int]) -> None:
    """Nur die genannten Zeilen behalten – Kopfzeile plus Musterzeile."""
    for i, tr in enumerate(list(rows(tbl_el))):
        if i not in keep:
            delete(tr)


# ---------------------------------------------------------------------------


def prepare(
    miete_path: str | Path,
    kauf_path: str | Path,
    out_path: str | Path,
) -> Path:
    """Ankerbasierte Vorlage aus den beiden Rohvorlagen erzeugen."""
    doc = docx.Document(str(miete_path))
    kauf = docx.Document(str(kauf_path))
    loc = Locator(doc)

    _prepare_deckblatt(loc)
    _prepare_toc(loc)
    _prepare_kapitel_1(loc)
    _prepare_vertragstext(loc, kauf)
    _prepare_hardware(loc)
    _prepare_konditionen(loc)
    _prepare_schluss(loc)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    validate_template(out)
    return out


# --- Deckblatt -------------------------------------------------------------


def _prepare_deckblatt(loc: Locator) -> None:
    grids = [t for t in loc.tables() if _tbl_style(t) == "TableGrid"]
    if len(grids) < 5:
        raise OfferteError("E101", "Vorlage: Deckblatt-Tabellen fehlen")
    t_kunde, t_anbieter, t_kopf = grids[2], grids[3], grids[4]

    # Kundenadressblock
    wrap_in_sdt(_cell_blocks(_cell(t_kunde, 0, 2)), "OFF.KUNDE")

    # Anbieteradresse (statisch) und Ansprechperson
    zelle = _cell(t_anbieter, 0, 2)
    blocks = _cell_blocks(zelle)
    # Die Beschriftungen in Spalte 1 stehen auf festen Zeilen: "Anbieter" auf
    # Zeile 1, "Ihre Ansprechperson" auf Zeile 5.  Der Leerabsatz nach der
    # Adresse gehört deshalb zum statischen Anbieterblock – sonst rutscht der
    # Kontakt eine Zeile hoch und die Beschriftung steht neben der falschen.
    wrap_in_sdt(blocks[0:4], "OFF.ANBIETER")
    wrap_in_sdt(blocks[4:8], "OFF.KONTAKT")
    # Die leere Vorlagenzeile "Direkt\t\t\t\t" entfällt: die Direktwahl wird als
    # Zeile des Blocks OFF.KONTAKT gesetzt, sonst stünde sie zweimal im Dokument.
    for b in blocks[8:]:
        if paragraph_text(b).strip().startswith("Direkt"):
            delete(b)

    # Klassifizierung / Nummer / Version / Datum: die Vorlage bringt bereits
    # drei Steuerelemente mit, sie werden nur benannt.
    zelle = _cell(t_kopf, 0, 2)
    vorhandene = [b for b in _cell_blocks(zelle) if b.tag == W("w:sdt")]
    if len(vorhandene) < 3:
        raise OfferteError("E101", "Vorlage: Deckblatt-Steuerelemente fehlen")
    for sdt, tag in zip(vorhandene, ("OFF.KLASSIFIZIERUNG", "OFF.NUMMER", "OFF.VERSION")):
        _retag(sdt, tag)
    datum_p = [b for b in _cell_blocks(zelle) if b.tag == W("w:p")]
    wrap_in_sdt(datum_p[-1:], "OFF.DATUM")

    wrap_in_sdt(_cell_blocks(_cell(t_kopf, 1, 2)), "OFF.GUELTIG_BIS")


def _retag(sdt_el, tag: str) -> None:
    pr = sdt_el.find(W("w:sdtPr"))
    if pr is None:
        pr = sdt_el.makeelement(W("w:sdtPr"), {})
        sdt_el.insert(0, pr)
    t = pr.find(W("w:tag"))
    if t is None:
        t = pr.makeelement(W("w:tag"), {})
        pr.append(t)
    t.set(qn("w:val"), tag)
    alias = pr.find(W("w:alias"))
    if alias is None:
        alias = pr.makeelement(W("w:alias"), {})
        pr.insert(0, alias)
    alias.set(qn("w:val"), tag)


def _prepare_toc(loc: Locator) -> None:
    """Das Inhaltsverzeichnis der Vorlage benennen (Abschnitt 12.1)."""
    for sdt in loc.body.findall(W("w:sdt")):
        text = paragraph_text(sdt)
        if text.strip().startswith("Inhalt"):
            _retag(sdt, "SYS.TOC")
            return
    raise OfferteError("E101", "Vorlage: Inhaltsverzeichnis nicht gefunden")


# --- Kapitel 1 -------------------------------------------------------------


def _prepare_kapitel_1(loc: Locator) -> None:
    head_standort = loc.para("Heading3", "Standort", 0)
    tbl_hardware = loc.table_by_style("graphax11")
    tbl_total = loc.table_by_style("graphax100")
    head_dl = loc.para("Heading3", "Im Angebot enthaltene")
    head_service_standort = loc.para("Heading3", "Standort", 1)
    tbl_service = loc.table_by_style("graphax1000")
    head_vertrag = loc.para("Heading2", "Laufzeit und Kündigungsfrist")

    # --- Hardwaretabelle: Kopfzeile + eine Musterzeile -------------------
    _keep_rows(tbl_hardware, [0, 1])
    set_tbl_look(tbl_hardware, "04A0")  # ohne lastRow (Abschnitt 10.4)

    # --- Adresszeile unter der Standortüberschrift -----------------------
    adress_p = make_paragraph(head_standort, "Installationsadresse", "05Klein")
    head_standort.addnext(adress_p)

    # --- Dienstleistungstabelle aus der Summentabelle klonen -------------
    # Die Vorlage kennt keine solche Tabelle; ein Neuaufbau verlöre tblGrid
    # und Rahmen, deshalb wird die vorhandene graphax100-Tabelle geklont.
    tbl_dl = clone(tbl_total)
    _keep_rows(tbl_dl, [0])
    muster = clone(rows(tbl_dl)[0])
    rows(tbl_dl)[0].addnext(muster)
    set_cell_paragraphs(cells(rows(tbl_dl)[0])[0], ["Leistung"])
    set_cell_paragraphs(cells(rows(tbl_dl)[0])[1], ["Betrag"])
    set_tbl_look(tbl_dl, "04E0")  # mit Summenzeile
    head_dl.addnext(tbl_dl)

    # --- Servicetabelle: Kopfzeile + eine Musterzeile --------------------
    _keep_rows(tbl_service, [0, 1])
    set_tbl_look(tbl_service, "04E0")

    # --- Kapitel 1.3 "Total" herstellen (Abschnitt 9) --------------------
    # Die Rohvorlage stellt die Summe direkt unter die Geräteliste.  Abschnitt 9
    # verlangt ein eigenes Kapitel nach den Servicepreisen.
    head_total = make_paragraph(head_vertrag, "Total", "Heading2")
    head_vertrag.addprevious(head_total)
    delete(tbl_total)
    _keep_rows(tbl_total, [0])
    set_tbl_look(tbl_total, "04E0")
    head_total.addnext(tbl_total)

    tbl_gesamt = clone(tbl_total)
    tbl_total.addnext(tbl_gesamt)

    # --- Anker setzen ----------------------------------------------------
    wrap_in_sdt([head_standort], "HEAD.STANDORT")
    wrap_in_sdt([adress_p], "LINE.STANDORT_ADRESSE")
    wrap_in_sdt([tbl_hardware], "TBL.HARDWARE")
    wrap_in_sdt([head_dl], "HEAD.DL")
    wrap_in_sdt([tbl_dl], "TBL.DIENSTLEISTUNG")
    wrap_in_sdt([head_service_standort], "HEAD.SERVICE_STANDORT")
    wrap_in_sdt([tbl_service], "TBL.SERVICE")
    wrap_in_sdt([tbl_total], "TBL.TOTAL")
    wrap_in_sdt([tbl_gesamt], "TBL.GESAMTTOTAL")
    wrap_in_sdt([head_vertrag], "HEAD.VERTRAGSTEXT")

    # Abschnittsanker umfassen die jeweils zusammengehörenden Anker.
    tm = tag_map(loc.body)
    sec_standort = [tm[t][0] for t in ("HEAD.STANDORT", "LINE.STANDORT_ADRESSE", "TBL.HARDWARE", "HEAD.DL", "TBL.DIENSTLEISTUNG")]
    wrap_in_sdt(sec_standort, "SEC.STANDORT")
    tm = tag_map(loc.body)
    sec_service = [tm[t][0] for t in ("HEAD.SERVICE_STANDORT", "TBL.SERVICE")]
    wrap_in_sdt(sec_service, "SEC.SERVICE")


# --- Vertragstext ----------------------------------------------------------


def _prepare_vertragstext(loc: Locator, kauf_doc) -> None:
    """Die drei Varianten zu einem SWITCH-Block zusammenfassen (Abschnitt 8.3)."""
    tm = tag_map(loc.body)
    head_sdt = tm["HEAD.VERTRAGSTEXT"][0]

    # Absätze zwischen Überschrift und der nächsten Überschrift sind der Miettext.
    # Lesezeichenmarken (``w:bookmarkStart``/``w:bookmarkEnd``) stehen dazwischen
    # und werden übersprungen.
    miete_paras = []
    node = head_sdt.getnext()
    while node is not None:
        if node.tag in (W("w:bookmarkStart"), W("w:bookmarkEnd")):
            node = node.getnext()
            continue
        if node.tag != W("w:p") or _p_style(node) in ("Heading1", "Heading2"):
            break
        miete_paras.append(node)
        node = node.getnext()
    if not miete_paras:
        raise OfferteError("E101", "Vorlage: Vertragstext (Miete) nicht gefunden")

    kauf_paras = _kauf_absaetze(kauf_doc, miete_paras[0])

    switch = make_sdt("SW.VERTRAGSTEXT")
    miete_paras[0].addprevious(switch)
    content = sdt_content(switch)

    for tag, quelle in (
        ("SW.VERTRAGSTEXT.KAUF", kauf_paras),
        ("SW.VERTRAGSTEXT.MIETE", miete_paras),
        ("SW.VERTRAGSTEXT.LEASING", [clone(p) for p in miete_paras]),
    ):
        var = make_sdt(tag)
        content.append(var)
        var_content = sdt_content(var)
        for p in quelle:
            delete(p)
            var_content.append(p)


def _kauf_absaetze(kauf_doc, muster_p) -> list:
    """Vertragstext der Kauf-Vorlage übernehmen."""
    body = kauf_doc.element.body
    treffer = []
    gefunden = False
    for child in body.iterchildren():
        if child.tag != W("w:p"):
            if gefunden and child.tag == W("w:tbl"):
                break
            continue
        style = _p_style(child)
        text = paragraph_text(child).strip()
        if style == "Heading2" and text.startswith("Laufzeit und Kündigungsfrist"):
            gefunden = True
            continue
        if gefunden:
            if style in ("Heading1", "Heading2"):
                break
            if paragraph_text(child).strip():
                treffer.append(clone(child))
    if not treffer:
        raise OfferteError("E101", "Kauf-Vorlage: Vertragstext nicht gefunden")
    return treffer


# --- Konditionen -----------------------------------------------------------


def _prepare_hardware(loc: Locator) -> None:
    """Platz für die Gerätedatenblätter schaffen (Kapitel nach den Preisen).

    Der Anker steht vor dem Kapitel «Konditionen». Bleibt er leer, wird er
    beim Rendern ersatzlos entfernt – ohne Datenblätter entsteht kein
    leeres Kapitel.
    """
    konditionen = loc.para("Heading1", "Konditionen")
    platzhalter = make_paragraph(konditionen, "", "Normal")
    konditionen.addprevious(platzhalter)
    wrap_in_sdt([platzhalter], "SEC.HARDWARE")


def _prepare_konditionen(loc: Locator) -> None:
    tbl = loc.table_by_style("graphax20")
    wrap_in_sdt([tbl], "TBL.KONDITIONEN")

    treffer = {}
    for tr in rows(tbl):
        tcs = cells(tr)
        if len(tcs) < 2:
            continue
        label = " ".join(paragraph_text(p) for p in tcs[0].findall(W("w:p"))).strip()
        if label.startswith("Abrechnungsintervall"):
            treffer["KOND.ABRECHNUNG"] = tcs[1]
        elif label.startswith("Rechnungsstellung"):
            treffer["KOND.RECHNUNG"] = tcs[1]
    for tag in ("KOND.ABRECHNUNG", "KOND.RECHNUNG"):
        if tag not in treffer:
            raise OfferteError("E101", f"Konditionentabelle: Zeile für {tag} fehlt")
        wrap_in_sdt(_cell_blocks(treffer[tag]), tag)


# --- Schlussteil -----------------------------------------------------------


def _prepare_schluss(loc: Locator) -> None:
    ort_datum = None
    for p in loc.paragraphs():
        if paragraph_text(p).strip().startswith("Spreitenbach,"):
            ort_datum = p
    if ort_datum is None:
        raise OfferteError("E101", "Vorlage: Zeile 'Spreitenbach, …' nicht gefunden")
    wrap_in_sdt([ort_datum], "OFF.ORT_DATUM")

    letzter = loc.blocks()[-1]
    nachweis = make_paragraph(ort_datum, "Kalkulationsgrundlage", "05Klein")
    letzter.addnext(nachweis)
    wrap_in_sdt([nachweis], "LINE.NACHWEIS")


# --- Prüfung ---------------------------------------------------------------


def validate_template(path: str | Path) -> dict[str, list]:
    """Ankerkatalog gegen die Vorlage halten (Abschnitt 13.1, E101/E102)."""
    doc = docx.Document(str(path))
    tm = tag_map(doc.element.body)
    vorhanden = set(tm)

    fehlend = [a.tag for a in BY_TAG.values() if not a.optional and a.tag not in vorhanden]
    if fehlend:
        raise OfferteError("E101", ", ".join(sorted(fehlend)))

    fehlende_varianten = [v for v in SWITCH_VARIANTS if v not in vorhanden]
    if fehlende_varianten:
        raise OfferteError("E101", ", ".join(fehlende_varianten))

    unbekannt = sorted(
        t for t in vorhanden if t not in ALL_TAGS and t != "Seitenumbruch nicht löschen"
    )
    if unbekannt:
        raise OfferteError("E102", ", ".join(unbekannt))
    return tm
