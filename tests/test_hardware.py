"""Gerätedatenblätter als Kapitel der Offerte.

Die Datenblätter sind Word-Vorlagen mit eigenen Formatvorlagen, Bildern und
Nummerierungen. Diese Tests sichern, dass die Übernahme alle drei umschlüsselt
und dass eine fehlende Zuordnung eine Warnung ist, kein Abbruch.
"""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

import docx
import pytest
from docx.oxml.ns import qn

from offerttool.docxutil.xmlutil import W, paragraph_text
from offerttool.errors import WarningCollector
from offerttool.hardware import Bibliothek, normalisieren
from offerttool.pipeline import generiere

from .conftest import KALKTOOL, VORLAGE

# Die Datenblätter sind Geschäftsunterlagen und liegen nicht im Repository.
DATENBLAETTER = Path(__file__).resolve().parent.parent / "datenblaetter"
ohne_datenblaetter = pytest.mark.skipif(
    not DATENBLAETTER.is_dir() or not any(DATENBLAETTER.rglob("*.dotx")),
    reason="Verzeichnis datenblaetter/ ist nicht vorhanden",
)


def stil(p) -> str:
    ppr = p.find(W("w:pPr"))
    if ppr is None:
        return "Normal"
    s = ppr.find(W("w:pStyle"))
    return s.get(qn("w:val")) if s is not None else "Normal"


def test_normalisieren_ignoriert_sprache_und_trennzeichen():
    assert normalisieren("bizhub C3351i de") == normalisieren("bizhub C3351i")
    assert normalisieren("brother MFC-L3750CDW de v2") == "brothermfcl3750cdw"


def test_bibliothek_ohne_verzeichnis_ist_leer(tmp_path):
    assert len(Bibliothek.laden(tmp_path)) == 0


def test_zuordnung_meldet_statt_zu_raten(tmp_path):
    """Passen mehrere Datenblätter, wird keines gewählt."""
    for name in ("bizhub C3351i de.dotx", "bizhub C3351i fr.dotx"):
        (tmp_path / name).write_bytes(b"")
    b = Bibliothek.laden(tmp_path)
    warn = WarningCollector()
    assert b.finde("bizhub C3351i", warn) is None
    assert "W330" in warn.codes()


def test_fehlendes_datenblatt_ist_eine_warnung(tmp_path):
    b = Bibliothek.laden(tmp_path)
    warn = WarningCollector()
    assert b.finde("bizhub C3351i", warn) is None
    assert "W331" in warn.codes()


@ohne_datenblaetter
def test_bibliothek_findet_das_geraet():
    b = Bibliothek.laden(DATENBLAETTER)
    warn = WarningCollector()
    treffer = b.finde("bizhub C3351i", warn)
    assert treffer is not None
    assert "C3351i" in treffer.modell
    assert warn.codes() == []


@ohne_datenblaetter
def test_kapitel_erscheint_mit_ueberschriften(tmp_path):
    ergebnis = generiere(
        [KALKTOOL], VORLAGE, tmp_path / "hw.docx",
        toc_seitenzahlen=False, datenblaetter_pfad=DATENBLAETTER,
    )
    d = docx.Document(str(ergebnis.offerte))
    ueberschriften = [
        (stil(p), paragraph_text(p).strip())
        for p in d.element.body.iter(W("w:p"))
        if stil(p) in ("Heading1", "Heading2", "Heading3")
    ]
    texte = [t for _s, t in ueberschriften]
    assert "Hardware" in texte
    assert "Multifunktionsgeräte" in texte
    # Die Modellüberschrift des Datenblatts muss als Heading3 ankommen, sonst
    # fehlte sie im Inhaltsverzeichnis.
    assert any(s == "Heading3" and "C3351i" in t for s, t in ueberschriften)


@ohne_datenblaetter
def test_bild_und_formatvorlagen_werden_uebernommen(tmp_path):
    ergebnis = generiere(
        [KALKTOOL], VORLAGE, tmp_path / "hw.docx",
        toc_seitenzahlen=False, datenblaetter_pfad=DATENBLAETTER,
    )
    with zipfile.ZipFile(ergebnis.offerte) as z:
        namen = z.namelist()
        medien = [n for n in namen if n.startswith("word/media/")]
        stile = z.read("word/styles.xml").decode("utf-8")
    # Das Gerätefoto muss mitgekommen sein.
    assert medien, "kein Bild im Ergebnis"
    # Die Formatvorlage der Optionsliste stammt aus dem Datenblatt.
    assert "Fliesstext10ptSpezifikationGerte" in stile


@ohne_datenblaetter
def test_schalter_laesst_die_spezifikation_weg(tmp_path):
    def text_von(mit: bool) -> str:
        ergebnis = generiere(
            [KALKTOOL], VORLAGE, tmp_path / f"hw_{mit}.docx",
            toc_seitenzahlen=False, datenblaetter_pfad=DATENBLAETTER,
            mit_spezifikation=mit,
        )
        d = docx.Document(str(ergebnis.offerte))
        return "\n".join(paragraph_text(p) for p in d.element.body.iter(W("w:p")))

    mit, ohne = text_von(True), text_von(False)
    assert "Artikel-Nr." in mit
    assert "Artikel-Nr." not in ohne
    # Beschreibung und Technikdaten bleiben in beiden Fassungen.
    assert "Die Vorteile auf einen Blick" in ohne
    assert len(ohne) < len(mit)


@ohne_datenblaetter
def test_dasselbe_geraet_erscheint_nur_einmal(tmp_path):
    import openpyxl

    zweit = tmp_path / "zweiter_standort.xlsx"
    shutil.copy(KALKTOOL, zweit)
    wb = openpyxl.load_workbook(zweit, data_only=True)
    wb.worksheets[0]["B7"] = "Werkhof"
    wb.save(zweit)

    ergebnis = generiere(
        [KALKTOOL, zweit], VORLAGE, tmp_path / "zwei.docx",
        toc_seitenzahlen=False, datenblaetter_pfad=DATENBLAETTER,
    )
    d = docx.Document(str(ergebnis.offerte))
    modelle = [
        paragraph_text(p).strip()
        for p in d.element.body.iter(W("w:p"))
        if stil(p) == "Heading3" and "C3351i" in paragraph_text(p)
    ]
    assert len(modelle) == 1, modelle


def test_ohne_verzeichnis_entsteht_kein_leeres_kapitel(tmp_path):
    """Ohne Datenblätter bleibt die Offerte unverändert – kein leerer Abschnitt."""
    ergebnis = generiere(
        [KALKTOOL], VORLAGE, tmp_path / "ohne.docx", toc_seitenzahlen=False
    )
    d = docx.Document(str(ergebnis.offerte))
    texte = [paragraph_text(p).strip() for p in d.element.body.iter(W("w:p"))]
    assert "Multifunktionsgeräte" not in texte
    assert "Konditionen" in texte
