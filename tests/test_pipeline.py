"""Pipeline von Ende zu Ende (Abschnitt 1) und die Abbruchregeln (13.1)."""

from __future__ import annotations

import shutil

import openpyxl
import pytest

from offerttool.errors import OfferteError
from offerttool.pipeline import generiere

from .conftest import CRM_JSON, KALKTOOL, VORLAGE


def dokumenttext(pfad) -> str:
    import docx

    from offerttool.docxutil.xmlutil import W, paragraph_text

    d = docx.Document(str(pfad))
    return "\n".join(paragraph_text(p) for p in d.element.body.iter(W("w:p")))


@pytest.fixture(scope="module")
def offerte(tmp_path_factory):
    ziel = tmp_path_factory.mktemp("out") / "Offerte.docx"
    return generiere([KALKTOOL], VORLAGE, ziel, CRM_JSON, toc_seitenzahlen=False)


def test_dateien_entstehen(offerte):
    assert offerte.offerte.exists()
    assert offerte.protokoll.exists()
    assert offerte.protokoll.name.endswith(".pruefprotokoll.md")


def test_kernwerte_im_dokument(offerte):
    text = dokumenttext(offerte.offerte)
    for erwartet in (
        "Gemeindeverwaltung Birsfelden",
        "Standort 1: Museum",
        "Installationsadresse: Schulstrasse 29 1.OG, 4127 Birsfelden",
        "bizhub C3351i",
        "CHF 0.0320",
        "CHF 0.0050",
        "CHF 3.50",
        "CHF 43.50",
        "CHF 47.00",
        "CHF 300.00",
        "28.07.2026",
        "26.09.2026",
        "01.08.2026",
    ):
        assert erwartet in text, erwartet


def test_gesperrte_werte_fehlen(offerte):
    """Abschnitt 14: insbesondere 2’645, 1’587, 2’024.75, 2’339.75 fehlen."""
    text = dokumenttext(offerte.offerte)
    for verboten in ("2’645", "1’587", "23.45", "2’024.75", "2’339.75"):
        assert verboten not in text, verboten


def test_keine_restplatzhalter(offerte):
    text = dokumenttext(offerte.offerte)
    assert "%%" not in text
    assert "{" not in text


def test_kein_vorlagentext_ueberlebt(offerte):
    """Steuerelemente in Zellen werden aufgelöst (Abschnitt 10.3)."""
    text = dokumenttext(offerte.offerte)
    for rest in ("bizhub C257i", "CHF 5.00", "Standort A", "36 Monaten"):
        assert rest not in text, rest


def test_kaufvariante_fehlt_bei_miete(offerte):
    text = dokumenttext(offerte.offerte)
    assert "Der Mietvertrag tritt am 01.08.2026 in Kraft" in text
    assert "beim Kauf" not in text


def test_gesamttotal_entfaellt_bei_einem_standort(offerte):
    text = dokumenttext(offerte.offerte)
    assert "alle Standorte" not in text


def test_protokoll_nennt_warnungen(offerte):
    inhalt = offerte.protokoll.read_text(encoding="utf-8")
    assert "W308" in inhalt
    assert "W312" in inhalt
    assert "MIETE" in inhalt


def test_determinismus(tmp_path):
    """Gleiche Eingabe -> gleiche Ausgabe (Grundsatz der Spezifikation)."""
    a = generiere([KALKTOOL], VORLAGE, tmp_path / "a.docx", CRM_JSON, toc_seitenzahlen=False)
    b = generiere([KALKTOOL], VORLAGE, tmp_path / "b.docx", CRM_JSON, toc_seitenzahlen=False)
    assert dokumenttext(a.offerte) == dokumenttext(b.offerte)


def _variiere(quelle, ziel, **zellen):
    """Abwandlung des Referenz-Kalktools als Testvorlage.

    ``data_only=True`` friert die zwischengespeicherten Formelergebnisse als
    Literale ein – sonst ginge beim Speichern der Wert jeder Formelzelle
    verloren und schon ``KM!M9`` wäre leer.
    """
    shutil.copy(quelle, ziel)
    wb = openpyxl.load_workbook(ziel, data_only=True)
    ws = wb.worksheets[0]
    for adresse, wert in zellen.items():
        ws[adresse] = wert
    wb.save(ziel)
    return ziel


def test_zwei_standorte(tmp_path):
    zweit = _variiere(KALKTOOL, tmp_path / "Kalktool_Standort2.xlsx", B7="Werkhof")

    ergebnis = generiere(
        [KALKTOOL, zweit], VORLAGE, tmp_path / "zwei.docx", CRM_JSON, toc_seitenzahlen=False
    )
    text = dokumenttext(ergebnis.offerte)
    assert "Standort 1: Museum" in text
    assert "Standort 2: Werkhof" in text
    # Bei mehr als einem Standort erscheint das Gesamttotal (Abschnitt 9).
    assert "Monatspauschale total – alle Standorte" in text
    assert "CHF 94.00" in text  # 2 x 47.00
    assert "CHF 600.00" in text  # 2 x 300.00


def test_kaufvariante(tmp_path):
    kauf = _variiere(KALKTOOL, tmp_path / "kauf.xlsx", I16=1)
    ergebnis = generiere(
        [kauf], VORLAGE, tmp_path / "kauf.docx", CRM_JSON, toc_seitenzahlen=False
    )
    text = dokumenttext(ergebnis.offerte)
    assert "Total Kauf" in text
    assert "CHF 2’339.75" in text  # C62, bei KAUF legitim
    assert "Laufzeit und Kündigungsfrist für Serviceverträge beim Kauf" in text
    assert "Mietpauschale" not in text


def test_e401_bei_unbekannter_finanzierungsart(tmp_path):
    datei = _variiere(KALKTOOL, tmp_path / "e401.xlsx", I16=9)
    with pytest.raises(OfferteError) as exc:
        generiere([datei], VORLAGE, tmp_path / "x.docx", toc_seitenzahlen=False)
    assert exc.value.code == "E401"


def test_e402_bei_verletzter_summenregel(tmp_path):
    datei = _variiere(KALKTOOL, tmp_path / "e402.xlsx", L95=99)
    with pytest.raises(OfferteError) as exc:
        generiere([datei], VORLAGE, tmp_path / "x.docx", toc_seitenzahlen=False)
    assert exc.value.code == "E402"


def test_e413_bei_miete_ohne_pauschale(tmp_path):
    datei = _variiere(KALKTOOL, tmp_path / "e413.xlsx", L92=0, L95=3.5)
    with pytest.raises(OfferteError) as exc:
        generiere([datei], VORLAGE, tmp_path / "x.docx", toc_seitenzahlen=False)
    assert exc.value.code == "E413"


def test_e403_bei_gemischten_varianten(tmp_path):
    zweit = _variiere(KALKTOOL, tmp_path / "e403.xlsx", I16=1)
    with pytest.raises(OfferteError) as exc:
        generiere([KALKTOOL, zweit], VORLAGE, tmp_path / "x.docx", toc_seitenzahlen=False)
    assert exc.value.code == "E403"


def test_e404_bei_verschiedenen_kunden(tmp_path):
    zweit = _variiere(KALKTOOL, tmp_path / "e404.xlsx", B5="Andere Firma AG")
    with pytest.raises(OfferteError) as exc:
        generiere([KALKTOOL, zweit], VORLAGE, tmp_path / "x.docx", toc_seitenzahlen=False)
    assert exc.value.code == "E404"


def test_keine_datei_bei_abbruch(tmp_path):
    """Bricht ein Schritt ab, entsteht keine Datei – nie eine halbe Offerte."""
    datei = _variiere(KALKTOOL, tmp_path / "abbruch.xlsx", I16=9)
    ziel = tmp_path / "nichts.docx"
    with pytest.raises(OfferteError):
        generiere([datei], VORLAGE, ziel, toc_seitenzahlen=False)
    assert not ziel.exists()


# --- Erscheinungsbild (nach Sichtprüfung des gerenderten PDF) --------------


def _tabellen(pfad):
    """Alle Tabellen des Dokuments mit Style, tblLook und Zeileninhalten."""
    import docx
    from docx.oxml.ns import qn

    from offerttool.docxutil.tables import cells, rows
    from offerttool.docxutil.xmlutil import W, paragraph_text

    d = docx.Document(str(pfad))
    out = []
    for tbl in d.element.body.iter(W("w:tbl")):
        pr = tbl.find(W("w:tblPr"))
        stil = pr.find(W("w:tblStyle")) if pr is not None else None
        look = pr.find(W("w:tblLook")) if pr is not None else None
        out.append({
            "stil": stil.get(qn("w:val")) if stil is not None else None,
            "look": look.get(qn("w:val")) if look is not None else None,
            "zeilen": [
                [" ".join(paragraph_text(p) for p in tc.findall(W("w:p"))).strip()
                 for tc in cells(tr)]
                for tr in rows(tbl)
            ],
        })
    return out


def test_servicetabelle_ohne_summenzeile(offerte):
    """Kapitel 1.2 hat keine Summenzeile, also darf lastRow nicht gesetzt sein.

    Mit lastRow setzt Word die letzte Zeile fett; die Zählerstandszeile läse
    sich dann wie ein Total (Abschnitt 5.4 und 10.4).
    """
    service = [t for t in _tabellen(offerte.offerte) if t["stil"] == "graphax1000"]
    assert service, "Servicetabelle nicht gefunden"
    for t in service:
        assert t["look"] == "04A0", t["look"]


def test_hardwaretabelle_ohne_leere_artikelspalte(offerte):
    """Führt das Kalktool keine Artikelnummern, entfällt die Spalte."""
    hardware = [t for t in _tabellen(offerte.offerte) if t["stil"] == "graphax11"]
    assert hardware, "Hardwaretabelle nicht gefunden"
    kopf = hardware[0]["zeilen"][0]
    assert kopf == ["Bezeichnung", "Stück"], kopf
    for zeile in hardware[0]["zeilen"]:
        assert len(zeile) == 2, zeile
        assert "–" not in zeile


def test_hardwaretabelle_zeigt_artikelnummern_wenn_vorhanden(tmp_path):
    """Sobald das Kalktool Artikelnummern liefert, erscheint die Spalte wieder."""

    from offerttool.derive import derive
    from offerttool.errors import WarningCollector
    from offerttool.extract import extract
    from offerttool.mapping import load_mapping
    from offerttool.workbook import Kalktool

    from .conftest import MAPPING

    # Die Zuordnung kennt heute keine Artikelnummernspalte; geprüft wird
    # deshalb der Renderer selbst, mit einer Position, die eine Nummer trägt.
    m = load_mapping(MAPPING)
    w = WarningCollector()
    with Kalktool(KALKTOOL, m, w) as kt:
        ctx = extract(kt, m, w)
    ctx.index = 1
    d = derive(ctx, w)
    ctx.listen["hardware"][0].artnr = "ACVD021"

    import docx

    from offerttool.crm import CRM
    from offerttool.derive import gesamttotal
    from offerttool.render import render

    doc = docx.Document(str(VORLAGE))
    render(doc, [(ctx, d)], gesamttotal([(ctx, d)]), m, CRM({}), w)
    ziel = tmp_path / "mit_artnr.docx"
    doc.save(str(ziel))

    hardware = [t for t in _tabellen(ziel) if t["stil"] == "graphax11"]
    assert hardware[0]["zeilen"][0] == ["Artikel No.", "Bezeichnung", "Stück"]
    assert hardware[0]["zeilen"][1][0] == "ACVD021"


def test_anbieterblock_haelt_die_beschriftungen_auf_hoehe(offerte):
    """Die Beschriftung "Ihre Ansprechperson" steht auf Zeile 5 der Zelle.

    Der Leerabsatz nach der Graphax-Adresse gehört deshalb zum statischen
    Anbieterblock; ohne ihn rutscht der Kontakt hoch und die Beschriftung
    steht neben "Telefon Zentrale".
    """
    import docx

    from offerttool.docxutil.tables import cells, rows
    from offerttool.docxutil.xmlutil import W, paragraph_text

    d = docx.Document(str(offerte.offerte))
    for tbl in d.element.body.iter(W("w:tbl")):
        zeilen = rows(tbl)
        if not zeilen:
            continue
        tcs = cells(zeilen[0])
        if len(tcs) < 3:
            continue
        beschriftung = [paragraph_text(p).strip() for p in tcs[0].findall(W("w:p"))]
        if "Anbieter" not in beschriftung:
            continue
        inhalt = [paragraph_text(p).strip() for p in tcs[2].findall(W("w:p"))]
        zeile = beschriftung.index("Ihre Ansprechperson")
        assert inhalt[zeile] == "Thomas Steiner", (zeile, inhalt[:8])
        return
    raise AssertionError("Anbieterblock nicht gefunden")


# --- Sperrliste: Fehlalarm durch statischen Vorlagentext ------------------


def test_stundensatz_der_vorlage_ist_kein_leck(tmp_path):
    """Die Konditionentabelle nennt 180 CHF pro Stunde.

    Steht 180 zufällig auch in einer gesperrten Zelle des Kalktools, ist das
    kein Leck: die Zahl stand im Dokument, bevor überhaupt ein Kalktool
    gelesen wurde. Vorher brach der Generator hier mit E601 ab.
    """
    datei = _variiere(KALKTOOL, tmp_path / "mit180.xlsx", E53=180)
    ergebnis = generiere(
        [datei], VORLAGE, tmp_path / "mit180.docx", toc_seitenzahlen=False
    )
    text = dokumenttext(ergebnis.offerte)
    assert "180.- CHF pro Stunde" in text
    assert "Gemeindeverwaltung Birsfelden" in text


@pytest.mark.parametrize("wert", [120, 130, 200])
def test_weitere_zahlen_der_vorlage_ebenso(tmp_path, wert):
    """120, 130 und 200 stehen ebenfalls in den Konditionen."""
    datei = _variiere(KALKTOOL, tmp_path / f"w{wert}.xlsx", E53=wert)
    ergebnis = generiere(
        [datei], VORLAGE, tmp_path / f"w{wert}.docx", toc_seitenzahlen=False
    )
    assert ergebnis.offerte.exists()


def test_echte_sperrwerte_werden_weiterhin_erkannt():
    """Die Ausnahme darf die Prüfung nicht aushebeln."""
    from offerttool.errors import OfferteError
    from offerttool.validate import statische_token, validate_output

    statisch = statische_token(VORLAGE)
    assert "180" in statisch, "Stundensatz nicht als statisch erkannt"

    for wert in ("2’645", "1’587", "2’024.75"):
        assert wert not in statisch
        with pytest.raises(OfferteError) as exc:
            validate_output(f"Im Dokument steht {wert}", {wert}, statisch)
        assert exc.value.code == "E601"


def test_vorlagenreste_gelten_nicht_als_statisch():
    """Was in einem Anker steht, soll überschrieben werden.

    Bleibt eine Zahl von dort stehen, ist das ein Vorlagenrest und muss
    weiterhin auffallen – die Musterzeilen dürfen deshalb nicht in die
    Ausnahmeliste geraten.
    """
    from offerttool.validate import statische_token

    statisch = statische_token(VORLAGE)
    # Beträge aus den Musterzeilen der Vorlage
    for rest in ("112.65", "257"):
        assert rest not in statisch, f"{rest} wäre als Rest nicht mehr erkennbar"
