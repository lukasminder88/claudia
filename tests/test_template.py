"""Ankerkatalog und Vorlagenpräparation (Abschnitt 3 und 13.1)."""

from __future__ import annotations

import docx
import pytest

from offerttool.anchors import ALL_TAGS, BY_TAG, SWITCH_VARIANTS
from offerttool.docxutil.anchor_ops import tag_map
from offerttool.docxutil.xmlutil import W, sdt_tag
from offerttool.errors import OfferteError
from offerttool.prepare import prepare, validate_template

from .conftest import ROH_KAUF, ROH_MIETE, VORLAGE


def test_vorlage_erfuellt_den_katalog():
    tm = validate_template(VORLAGE)
    for anker in BY_TAG.values():
        assert anker.tag in tm, anker.tag
    for variante in SWITCH_VARIANTS:
        assert variante in tm


def test_jeder_anker_genau_einmal():
    tm = validate_template(VORLAGE)
    mehrfach = {t: len(v) for t, v in tm.items() if len(v) > 1 and t in ALL_TAGS}
    assert mehrfach == {}


def test_praeparation_ist_reproduzierbar(tmp_path):
    a = prepare(ROH_MIETE, ROH_KAUF, tmp_path / "a.docx")
    b = prepare(ROH_MIETE, ROH_KAUF, tmp_path / "b.docx")
    assert _struktur(a) == _struktur(b)


def test_praeparation_trifft_die_ausgelieferte_vorlage(tmp_path):
    neu = prepare(ROH_MIETE, ROH_KAUF, tmp_path / "neu.docx")
    assert _struktur(neu) == _struktur(VORLAGE)


def _struktur(pfad) -> list[str]:
    d = docx.Document(str(pfad))
    return [
        sdt_tag(s) or "?"
        for s in d.element.body.findall(".//" + W("w:sdt"))
    ]


def test_e101_bei_fehlendem_anker(tmp_path):
    """Eine Rohvorlage ohne Anker wird abgelehnt."""
    with pytest.raises(OfferteError) as exc:
        validate_template(ROH_MIETE)
    assert exc.value.code == "E101"


def test_e102_bei_unbekanntem_anker(tmp_path):
    d = docx.Document(str(VORLAGE))
    from offerttool.docxutil.xmlutil import make_sdt

    d.element.body.append(make_sdt("OFF.ERFUNDEN"))
    ziel = tmp_path / "unbekannt.docx"
    d.save(str(ziel))
    with pytest.raises(OfferteError) as exc:
        validate_template(ziel)
    assert exc.value.code == "E102"


def test_musterzeilen_sind_vorhanden():
    """Jede Tabelle hat Kopf- und Musterzeile (Abschnitt 10.1)."""
    d = docx.Document(str(VORLAGE))
    tm = tag_map(d.element.body)
    erwartet = {
        "TBL.HARDWARE": 2,
        "TBL.DIENSTLEISTUNG": 2,
        "TBL.SERVICE": 2,
        "TBL.TOTAL": 1,
        "TBL.GESAMTTOTAL": 1,
    }
    for tag, zeilen in erwartet.items():
        tbl = tm[tag][0].find(".//" + W("w:tbl"))
        assert tbl is not None, tag
        assert len(tbl.findall(W("w:tr"))) == zeilen, tag


def test_switch_hat_drei_varianten():
    d = docx.Document(str(VORLAGE))
    tm = tag_map(d.element.body)
    switch = tm["SW.VERTRAGSTEXT"][0]
    kinder = [sdt_tag(s) for s in switch.findall(".//" + W("w:sdt"))]
    assert set(SWITCH_VARIANTS) <= set(kinder)
